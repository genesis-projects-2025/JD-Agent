# backend/app/services/docx_generator.py
# Pulse Pharma branded DOCX — matches the official company PDF/JD template exactly.
# Compact 2-page layout, grey section headers (#BFBFBF), small Pulse logo + brand title in header,
# "About Pulse" section, exact table structure, and footer disclaimer.

from io import BytesIO
import logging
import os
from urllib.request import urlopen

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

_HERE = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(os.path.dirname(_HERE))

DEFAULT_LOCAL_LOGO = os.path.join(_HERE, "..", "static", "images", "logo.png")
DEFAULT_LOGO_URL = "https://company-logo-wtn.s3.ap-southeast-2.amazonaws.com/logo.png"

HEADER_COLOR = "BFBFBF"  # exact grey from company template
BORDER_COLOR = "999999"
BRAND_PURPLE = RGBColor(0x5B, 0x20, 0x53)  # Pulse brand color #5B2053

ABOUT_PULSE_TEXT = (
    "Pulse is a fast-growing Pharmaceutical company with a vertically & diagonally integrated business model, "
    "focused on providing innovative product solutions to a large number of people around the world, to help them "
    "manage their health better & lead a quality life. We are passionate for Innovation and compassionate for people. "
    "We go by the philosophy, solving the unsolved, reaching the unreached and serving the unserved.\n\n"
    "We believe that health and wellbeing are the main sources of happiness for humankind. Our goal is to preserve "
    "that happiness by developing and producing patient friendly medicines."
)

# ── Low-level OOXML helpers ───────────────────────────────────────────────────

def _set_cell_properties(cell, bg_color: str | None = None, borders: bool = True, valign: str | None = None) -> None:
    """Sets cell properties ensuring correct OOXML element order in tcPr."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 1. Borders
    if borders:
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is not None:
            tcPr.remove(tcBorders)
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), BORDER_COLOR)
            tcBorders.append(b)
        tcPr.append(tcBorders)

    # 2. Background Shading
    if bg_color:
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            tcPr.remove(shd)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_color)
        tcPr.append(shd)

    # 3. Vertical Alignment
    if valign:
        va = tcPr.find(qn("w:vAlign"))
        if va is not None:
            tcPr.remove(va)
        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), valign)
        tcPr.append(va)


def _para_spacing(para, before: float = 0, after: float = 2, line_spacing: float = 1.15) -> None:
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line_spacing


def _load_logo_stream() -> BytesIO | None:
    """Load company logo image: local static file first, then URL fallback."""
    local_path = os.path.abspath(DEFAULT_LOCAL_LOGO)
    if os.path.exists(local_path):
        try:
            with open(local_path, "rb") as f:
                return BytesIO(f.read())
        except Exception as exc:
            logger.warning("Failed loading local logo from %s: %s", local_path, exc)

    if DEFAULT_LOGO_URL:
        try:
            with urlopen(DEFAULT_LOGO_URL, timeout=8) as response:
                return BytesIO(response.read())
        except Exception as exc:
            logger.warning("Failed loading logo from URL %s: %s", DEFAULT_LOGO_URL, exc)

    return None

# ── Row builders ──────────────────────────────────────────────────────────────

def _section_header_row(table, row_idx: int, text: str) -> None:
    """Merge cols, grey background (#BFBFBF), 12pt bold centred text."""
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    _set_cell_properties(cell, bg_color=HEADER_COLOR, borders=True)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(para, before=4, after=4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _sub_header_row(table, row_idx: int, text: str) -> None:
    """Merge cols, grey background (#BFBFBF), 11pt bold centred text."""
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    _set_cell_properties(cell, bg_color=HEADER_COLOR, borders=True)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(para, before=3, after=3)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _data_row(table, row_idx: int, label: str, value) -> None:
    """Bold label cell | plain value cell."""
    row = table.rows[row_idx]
    lc, vc = row.cells[0], row.cells[1]

    for cell in (lc, vc):
        _set_cell_properties(cell, borders=True, valign="top")

    # Label
    lp = lc.paragraphs[0]
    _para_spacing(lp, before=2, after=2)
    lr = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(11)

    # Value
    if isinstance(value, list):
        if not value:
            vp = vc.paragraphs[0]
            _para_spacing(vp, before=2, after=2)
            vr = vp.add_run("To be confirmed with line manager.")
            vr.font.size = Pt(11)
        else:
            for i, item in enumerate(value):
                p = vc.paragraphs[0] if i == 0 else vc.add_paragraph()
                _para_spacing(p, before=1, after=2)
                p.paragraph_format.left_indent = Inches(0.15)
                r = p.add_run(f"\u2022 {item}")
                r.font.size = Pt(11)
    else:
        vp = vc.paragraphs[0]
        _para_spacing(vp, before=2, after=2)
        vr = vp.add_run(str(value) if value else "—")
        vr.font.size = Pt(11)

# ── Data extraction helpers ───────────────────────────────────────────────────

def _get(data: dict, *keys) -> str:
    emp = data.get("employee_information") or {}
    for k in keys:
        v = data.get(k) or emp.get(k)
        if not v and data.get("qualifications"):
            v = data["qualifications"].get(k)
        if v and isinstance(v, str) and v.strip():
            return v.strip()
    return ""


def _get_list(data: dict, *keys) -> list:
    for k in keys:
        v = data.get(k)
        if not v and data.get("qualifications"):
            v = data["qualifications"].get(k)
        if isinstance(v, list) and v:
            return [str(x).strip() for x in v if x and str(x).strip()]
        if isinstance(v, str) and v.strip():
            return [s.replace("-", "").replace("*", "").replace("•", "").strip() for s in v.split("\n") if s.strip()]
    return []

# ── Main Generator ─────────────────────────────────────────────────────────────

def generate_jd_docx(
    jd_data: dict,
    title: str | None = None,
    department: str | None = None,
    kra_kpi_data: dict | None = None,
) -> BytesIO:
    """
    Generate a compact, Pulse Pharma branded DOCX matching the official PDF template.
    Fits comfortably within ~2 pages.
    """
    designation = _get(jd_data, "job_title", "title", "designation") or title or "—"
    job_level = _get(jd_data, "job_level", "joblevel", "grade") or "—"
    func = _get(jd_data, "department", "function") or department or "—"
    location = _get(jd_data, "location") or "Head Office"
    reporting_to = (
        _get(jd_data, "reports_to", "reporting_to")
        or (jd_data.get("working_relationships") or {}).get("reporting_to")
        or (jd_data.get("team_structure") or {}).get("reports_to")
        or "—"
    )

    purpose = _get(jd_data, "purpose", "role_summary")
    responsibilities = _get_list(jd_data, "responsibilities", "key_responsibilities", "tasks", "priority_tasks")
    skills = _get_list(jd_data, "skills", "technical_skills", "required_skills")
    tools = _get_list(jd_data, "tools", "tools_used", "tools_and_technologies")

    education = _get(jd_data, "education", "educational_qualification")
    experience = _get(jd_data, "experience", "relevant_experience")
    edu_exp = "\n\n".join(filter(None, [education, experience])) or "To be confirmed with line manager."

    # ── Document Setup ────────────────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0]
    section.page_width = Emu(7556500)   # A4
    section.page_height = Emu(10680700)
    section.top_margin = Emu(457200)    # 0.5 inch margins for 2-page fit
    section.bottom_margin = Emu(457200)
    section.left_margin = Emu(548640)   # 0.6 inch margins
    section.right_margin = Emu(548640)

    normal_style = doc.styles["Normal"]
    if hasattr(normal_style, "font"):
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

    # ── Header Logo & Title (Small height, exact PDF branding) ─────────────────
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(hp, before=0, after=4)

    logo_stream = _load_logo_stream()
    if logo_stream:
        # Small neat logo height (0.45 inches / 32pt)
        run_logo = hp.add_run()
        run_logo.add_picture(logo_stream, height=Inches(0.45))
        
        # Add brand text "Pulse" next to logo in brand color
        run_text = hp.add_run("  Pulse")
        run_text.font.name = "Arial"
        run_text.font.size = Pt(24)
        run_text.font.bold = True
        run_text.font.color.rgb = BRAND_PURPLE
    else:
        run_text = hp.add_run("Pulse Pharma")
        run_text.font.name = "Arial"
        run_text.font.size = Pt(20)
        run_text.font.bold = True
        run_text.font.color.rgb = BRAND_PURPLE

    # ── TABLE 1: Job / Role Information & Job Description ──────────────────────
    t1 = doc.add_table(rows=8, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER

    _section_header_row(t1, 0, "Job / Role Information")
    _data_row(t1, 1, "Designation", designation)
    _data_row(t1, 2, "Job Level", job_level)
    _data_row(t1, 3, "Department", func)
    _data_row(t1, 4, "Location", location)
    _data_row(t1, 5, "Reporting Manager", reporting_to)

    # About Pulse Subheader & Content
    _sub_header_row(t1, 6, "About Pulse")
    row_about = t1.rows[7]
    row_about.cells[0].merge(row_about.cells[1])
    cell_about = row_about.cells[0]
    _set_cell_properties(cell_about, borders=True, valign="top")
    p_about = cell_about.paragraphs[0]
    _para_spacing(p_about, before=4, after=4)
    p_about.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_about = p_about.add_run(ABOUT_PULSE_TEXT)
    r_about.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── TABLE 2: Job Description ──────────────────────────────────────────────
    t_jd = doc.add_table(rows=2, cols=2)
    t_jd.alignment = WD_TABLE_ALIGNMENT.CENTER
    _sub_header_row(t_jd, 0, "Job Description")

    row_jd = t_jd.rows[1]
    row_jd.cells[0].merge(row_jd.cells[1])
    cell_jd = row_jd.cells[0]
    _set_cell_properties(cell_jd, borders=True, valign="top")

    # Purpose
    if purpose:
        pp = cell_jd.paragraphs[0]
        _para_spacing(pp, before=3, after=3)
        r_pur_lbl = pp.add_run("Purpose of the Job / Role :\n")
        r_pur_lbl.bold = True
        r_pur_lbl.font.size = Pt(11)
        r_pur_val = pp.add_run(purpose)
        r_pur_val.font.size = Pt(11)
    else:
        cell_jd.paragraphs[0].clear()

    # Responsibilities
    if responsibilities:
        pr = cell_jd.add_paragraph()
        _para_spacing(pr, before=4, after=3)
        r_resp_lbl = pr.add_run("Job Responsibilities")
        r_resp_lbl.bold = True
        r_resp_lbl.font.size = Pt(11)

        for resp in responsibilities:
            rp = cell_jd.add_paragraph()
            _para_spacing(rp, before=1, after=2)
            rp.paragraph_format.left_indent = Inches(0.18)
            r_item = rp.add_run(f"\u2022 {resp}")
            r_item.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── TABLE 3: Skills / Competencies Required ──────────────────────────────
    t3 = doc.add_table(rows=3, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER

    _section_header_row(t3, 0, "Skills/ Competencies Required")
    _data_row(t3, 1, "Skills", skills if skills else ["To be confirmed with line manager."])
    _data_row(t3, 2, "Tools / Platforms", tools if tools else ["To be confirmed with line manager."])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── TABLE 4: Academic Qualifications & Experience Required ──────────────
    t4 = doc.add_table(rows=2, cols=2)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER

    _section_header_row(t4, 0, "Academic Qualifications & Experience Required")
    row4 = t4.rows[1]
    lc4, vc4 = row4.cells[0], row4.cells[1]
    for cell in (lc4, vc4):
        _set_cell_properties(cell, borders=True, valign="top")

    lp4 = lc4.paragraphs[0]
    _para_spacing(lp4, before=2, after=2)
    lr4 = lp4.add_run("Required Educational Qualification &\nRelevant experience")
    lr4.bold = True
    lr4.font.size = Pt(11)

    vp4 = vc4.paragraphs[0]
    _para_spacing(vp4, before=2, after=2)
    vr4 = vp4.add_run(edu_exp)
    vr4.font.size = Pt(10.5)

    # ── Optional KRA / KPI Framework Section ──────────────────────────────────
    if kra_kpi_data and kra_kpi_data.get("kras"):
        doc.add_page_break()

        kp_title_para = doc.add_paragraph()
        kp_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(kp_title_para, before=8, after=4)
        kp_title_run = kp_title_para.add_run(
            "Key Result Areas (KRAs) & Key Performance Indicators (KPIs)"
        )
        kp_title_run.bold = True
        kp_title_run.font.size = Pt(13)

        kras = kra_kpi_data["kras"]
        for kra in kras:
            num_kpis = len(kra.get("kpis", []))
            total_rows = 3 + num_kpis

            t_kra = doc.add_table(rows=total_rows, cols=2)
            t_kra.alignment = WD_TABLE_ALIGNMENT.CENTER

            # KRA Header
            r0 = t_kra.rows[0]
            r0.cells[0].merge(r0.cells[1])
            _set_cell_properties(r0.cells[0], bg_color=HEADER_COLOR, borders=True)
            p0 = r0.cells[0].paragraphs[0]
            _para_spacing(p0, before=2, after=2)
            run0 = p0.add_run(f"KRA: {kra.get('title')} (Weight: {kra.get('weight', 0)}%)")
            run0.bold = True
            run0.font.size = Pt(11)

            # KRA Description
            r1 = t_kra.rows[1]
            r1.cells[0].merge(r1.cells[1])
            _set_cell_properties(r1.cells[0], borders=True, valign="top")
            p1 = r1.cells[0].paragraphs[0]
            _para_spacing(p1, before=2, after=2)
            p1.add_run("Description: ").bold = True
            run1 = p1.add_run(kra.get("description", ""))
            run1.font.size = Pt(10.5)

            # KPI Section Subheader
            r2 = t_kra.rows[2]
            r2.cells[0].merge(r2.cells[1])
            _set_cell_properties(r2.cells[0], bg_color="E6E6E6", borders=True)
            p2 = r2.cells[0].paragraphs[0]
            _para_spacing(p2, before=2, after=2)
            run2 = p2.add_run("Key Performance Indicators (KPIs)")
            run2.bold = True
            run2.font.size = Pt(10.5)

            # KPIs details
            for idx, kpi in enumerate(kra.get("kpis", [])):
                row_idx = 3 + idx
                row = t_kra.rows[row_idx]
                lc, rc = row.cells[0], row.cells[1]
                for cell in (lc, rc):
                    _set_cell_properties(cell, borders=True, valign="top")

                lp = lc.paragraphs[0]
                _para_spacing(lp, before=2, after=2)
                lmr = lp.add_run(f"{kpi.get('metric')}\n")
                lmr.bold = True
                lmr.font.size = Pt(10.5)

                if kpi.get("description"):
                    ldr = lp.add_run(kpi.get("description"))
                    ldr.font.size = Pt(9.5)
                    ldr.italic = True

                rp = rc.paragraphs[0]
                _para_spacing(rp, before=2, after=2)
                rp.add_run("Target: ").bold = True
                rp.add_run(f"{kpi.get('target')}\n").font.size = Pt(10)

                rp.add_run("Measured Via: ").bold = True
                rp.add_run(f"{kpi.get('measurement_method')}\n").font.size = Pt(10)

                rp.add_run("Frequency: ").bold = True
                rp.add_run(f"{kpi.get('frequency')}\n").font.size = Pt(10)

                thresh = kpi.get("threshold", {})
                if thresh:
                    rp.add_run("Thresholds:\n").bold = True
                    rp.add_run(f"  • Below: {thresh.get('below_expectation', '')}\n").font.size = Pt(9.5)
                    rp.add_run(f"  • Meets: {thresh.get('meets_expectation', '')}\n").font.size = Pt(9.5)
                    rp.add_run(f"  • Excellent: {thresh.get('excellent', '')}").font.size = Pt(9.5)

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Footer Disclaimer ─────────────────────────────────────────────────────
    fp = doc.add_paragraph()
    _para_spacing(fp, before=6, after=0)
    fr = fp.add_run(
        "Pulse Pharma is an equal opportunity employer - we never differentiate candidates "
        "on the basis of religion, caste, gender, language, disabilities or ethnic group. "
        "Pulse reserves the right to place/move any candidate to any company location, "
        "partner location or customer location globally, in the best interest of Pulse business."
    )
    fr.font.size = Pt(9)
    fr.font.italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
