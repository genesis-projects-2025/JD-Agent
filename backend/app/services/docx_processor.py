# backend/app/services/docx_processor.py
"""
DOCX Document Processor — Extract text from Word documents.

MATCHES PDFProcessor interface so both can be used interchangeably in jd_intelligence.py

Key methods (same as PDFProcessor):
- extract_text(docx_bytes) → str
- extract_metadata(docx_bytes) → Dict
- validate_docx(docx_bytes) → Tuple[bool, str]
"""

from __future__ import annotations

import io
import logging
from typing import Dict, Any, Tuple
from docx import Document

logger = logging.getLogger(__name__)


class DOCXProcessor:
    """Extract text and metadata from DOCX files (matches PDFProcessor interface)"""

    @staticmethod
    def extract_text(docx_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.

        INTERFACE MATCH: Same as PDFProcessor.extract_text()

        Args:
            docx_bytes: Raw DOCX file bytes

        Returns:
            Extracted text as string

        Raises:
            Exception: If DOCX extraction fails
        """
        try:
            # Load DOCX from bytes
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)

            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts)

            if not text.strip():
                raise Exception("No text could be extracted from DOCX")

            logger.info(f"Extracted {len(text)} characters from DOCX using python-docx")
            return text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")

    @staticmethod
    def extract_metadata(docx_bytes: bytes) -> Dict[str, Any]:
        """
        Extract DOCX metadata.

        INTERFACE MATCH: Same as PDFProcessor.extract_metadata()

        Args:
            docx_bytes: Raw DOCX file bytes

        Returns:
            Dictionary with metadata (same structure as PDF)
        """
        try:
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)

            # Count paragraphs and tables
            num_paragraphs = len(doc.paragraphs)
            num_tables = len(doc.tables)

            # Get core properties (metadata)
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
            }

            return {
                "num_pages": 1,  # DOCX doesn't have "pages" like PDF, so return 1
                "num_paragraphs": num_paragraphs,
                "num_tables": num_tables,
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Metadata extraction failed: {str(e)}")
            return {"num_pages": 0, "metadata": {}, "error": str(e)}

    @staticmethod
    def validate_docx(docx_bytes: bytes) -> Tuple[bool, str]:
        """
        Validate DOCX file.

        INTERFACE MATCH: Same as PDFProcessor.validate_pdf()

        Args:
            docx_bytes: Raw DOCX file bytes

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check minimum size
        if len(docx_bytes) < 1000:
            return False, "DOCX file is too small"

        # Check DOCX magic number (ZIP archive signature)
        # DOCX files are ZIP archives, so they start with 'PK'
        if not docx_bytes.startswith(b"PK"):
            return False, "Not a valid DOCX file (invalid ZIP signature)"

        # Try to load the document
        try:
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)

            # Check if document has content
            if not doc.paragraphs and not doc.tables:
                return False, "DOCX has no content"

        except Exception as e:
            return False, f"Invalid DOCX: {str(e)}"

        return True, "Valid DOCX"


# ── LEGACY SUPPORT: Aliases for consistency ──────────────────────────────────
# If code uses "extract_text_from_docx" instead of the class method,
# these standalone functions provide backward compatibility


async def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Standalone function (legacy). Use DOCXProcessor.extract_text() instead."""
    return DOCXProcessor.extract_text(docx_bytes)


async def extract_metadata_from_docx(docx_bytes: bytes) -> Dict[str, Any]:
    """Standalone function (legacy). Use DOCXProcessor.extract_metadata() instead."""
    return DOCXProcessor.extract_metadata(docx_bytes)


async def validate_docx_file(docx_bytes: bytes) -> Tuple[bool, str]:
    """Standalone function (legacy). Use DOCXProcessor.validate_docx() instead."""
    return DOCXProcessor.validate_docx(docx_bytes)