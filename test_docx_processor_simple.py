#!/usr/bin/env python3
"""
Test script to verify DOCX processing works - focusing on core functionality
without requiring full application configuration.
"""

import asyncio
import io
from docx import Document

# Add backend to path
import sys
sys.path.append('/Users/manideekshith/Desktop/JD-Agent/backend')

from app.services.docx_processor import DOCXProcessor


def create_sample_docx() -> bytes:
    """Create a sample JD DOCX for testing"""
    doc = Document()
    doc.add_paragraph('Senior Software Engineer')
    doc.add_paragraph('Department: Engineering')
    doc.add_paragraph('We are seeking a Senior Software Engineer to join our growing team.')
    doc.add_paragraph('')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('• Develop and maintain web applications using modern technologies')
    doc.add_paragraph('• Collaborate with product managers, designers, and other engineers')
    doc.add_paragraph('• Write clean, maintainable, and efficient code')
    doc.add_paragraph('• Participate in code reviews and contribute to team best practices')
    doc.add_paragraph('')
    doc.add_paragraph('Requirements:')
    doc.add_paragraph('• 5+ years of experience in software development')
    doc.add_paragraph('• Proficiency in Python, JavaScript, and React')
    doc.add_paragraph('• Experience with RESTful APIs and database design')
    doc.add_paragraph('• Bachelors degree in Computer Science or related field')
    doc.add_paragraph('• Strong problem-solving and communication skills')
    
    # Save to bytes
    docx_io = io.BytesIO()
    doc.save(docx_io)
    return docx_io.getvalue()


def test_docx_processor_directly():
    """Test DOCX processor directly without AI dependencies"""
    print("=" * 60)
    print("Testing DOCX Processor Core Functionality")
    print("=" * 60)
    
    # Create sample DOCX
    docx_bytes = create_sample_docx()
    print(f"✓ Created sample DOCX ({len(docx_bytes)} bytes)")
    
    # Test 1: Direct DOCX processor validation
    print("\n1. Testing DOCX Processor Validation...")
    is_valid, validation_msg = DOCXProcessor.validate_docx(docx_bytes)
    print(f"   Valid: {is_valid}")
    print(f"   Message: {validation_msg}")
    assert is_valid, f"DOCX validation failed: {validation_msg}"
    print("   ✓ DOCX validation PASSED")
    
    # Test 2: Text extraction
    print("\n2. Testing Text Extraction...")
    extracted_text = DOCXProcessor.extract_text(docx_bytes)
    print(f"   Extracted {len(extracted_text)} characters")
    print(f"   Preview: {extracted_text[:100]}...")
    assert len(extracted_text) > 50, "Extracted text too short"
    assert "Senior Software Engineer" in extracted_text, "Key content missing"
    print("   ✓ Text extraction PASSED")
    
    # Test 3: Metadata extraction
    print("\n3. Testing Metadata Extraction...")
    metadata = DOCXProcessor.extract_metadata(docx_bytes)
    print(f"   Metadata: {metadata}")
    assert metadata["num_pages"] >= 1, "Should have at least 1 page"
    assert "metadata" in metadata, "Metadata dictionary missing"
    print("   ✓ Metadata extraction PASSED")
    
    # Test 4: Interface matching with PDF processor
    print("\n4. Testing Interface Matching with PDF Processor...")
    from app.services.pdf_processor import PDFProcessor
    
    docx_methods = {m for m in dir(DOCXProcessor) if not m.startswith('_')}
    pdf_methods = {m for m in dir(PDFProcessor) if not m.startswith('_')}
    
    # Normalize method names for comparison
    docx_core = {m.replace('_docx', '') for m in docx_methods}
    pdf_core = {m.replace('_pdf', '') for m in pdf_methods}
    
    print(f"   DOCX core methods: {sorted(docx_core)}")
    print(f"   PDF core methods: {sorted(pdf_core)}")
    
    interfaces_match = docx_core == pdf_core
    print(f"   Interfaces match: {interfaces_match}")
    assert interfaces_match, "DOCX and PDF processors should have matching interfaces"
    print("   ✓ Interface matching PASSED")
    
    # Test 5: Verify the processor can handle edge cases
    print("\n5. Testing Edge Cases...")
    
    # Empty DOCX
    empty_doc = Document()
    empty_io = io.BytesIO()
    empty_doc.save(empty_io)
    empty_bytes = empty_io.getvalue()
    
    is_valid_empty, msg_empty = DOCXProcessor.validate_docx(empty_bytes)
    print(f"   Empty DOCX validation: {is_valid_empty} - {msg_empty}")
    # Empty DOCX should be invalid (no content)
    assert not is_valid_empty, "Empty DOCX should be invalid"
    print("   ✓ Empty DOCX handling PASSED")
    
    print("\n" + "=" * 60)
    print("ALL DOCX PROCESSOR TESTS PASSED!")
    print("The DOCX analyzer is working correctly and matches PDF functionality.")
    print("=" * 60)
    
    return True


if __name__ == "__main__":
    # Run the test
    try:
        success = test_docx_processor_directly()
        if success:
            print("\n🎉 DOCX processing verification successful!")
            exit(0)
        else:
            print("\n❌ DOCX processing verification failed!")
            exit(1)
    except Exception as e:
        print(f"\n💥 Test failed with exception: {e}")
        import traceback
        traceback.print_exc()
        exit(1)